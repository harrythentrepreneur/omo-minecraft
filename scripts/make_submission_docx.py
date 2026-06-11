"""Generate the Omo Mission Control hackathon submission as a formatted .docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── theme ────────────────────────────────────────────────────────────────────
INK = RGBColor(0x14, 0x14, 0x18)
CYAN = RGBColor(0x0E, 0x7C, 0x86)
GREY = RGBColor(0x55, 0x55, 0x5E)

doc = Document()

# base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = INK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.12


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = CYAN
    return p


def lead(text):
    """Bold opening claim line under a heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    return p


def body(runs):
    """runs: list of (text, bold) tuples for inline emphasis."""
    p = doc.add_paragraph()
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
    return p


def para(text):
    return body([(text, False)])


def bullet(runs):
    p = doc.add_paragraph(style="List Bullet")
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
    return p


def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("— — —")
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── title block ──────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.paragraph_format.space_after = Pt(2)
r = t.add_run("Omo Mission Control")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = INK

s = doc.add_paragraph()
s.paragraph_format.space_after = Pt(2)
r = s.add_run("A world that builds itself for autonomous AI teams.")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

m = doc.add_paragraph()
r = m.add_run("Google for Startups AI Agents Challenge · Track 1: Build (Net-New Agents)   ·   Gemini · ADK · MCP")
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

# ── 1. Problem to solve ──────────────────────────────────────────────────────
heading("Problem to solve")
lead("Teams are deploying fleets of AI agents — and flying blind.")
para(
    "You cannot see what an agent is doing. You cannot steer it once it's moving. And every new "
    "capability means a developer stops to pre-wire another agent, another integration, another "
    "dashboard. The “autonomous workforce” everyone was sold turns out to be a chat box bolted "
    "onto a frozen org-chart: the agents are invisible, the org is set in concrete, and the human is "
    "locked out of the loop exactly when judgment matters most."
)
para("Three failures, and they arrive together:")
bullet([("No observability", True), (" — the reasoning, the tool calls, the hand-offs between agents all happen in the dark.", False)])
bullet([("No steerability", True), (" — you can't redirect a task mid-flight; you wait for it to finish being wrong.", False)])
bullet([("No extensibility without engineering", True), (" — adding a function is a ticket, a sprint, a deploy — not a sentence you speak.", False)])
para(
    "This is the agent-ops gap. In 2026 it is acute, it is named, and it is unmet. The promise was an "
    "autonomous team. What shipped feels neither autonomous nor like a team."
)

rule()

# ── 2. Our solution ──────────────────────────────────────────────────────────
heading("Our solution")
lead("Omo turns an autonomous AI organisation into a place you walk through — and the place builds itself to fit the mission.")
para(
    "You drop a small, glowing HQ anywhere on an ordinary Minecraft map. Inside, three Gemini agents are "
    "already at their desks, their live reasoning streaming on the screens behind them. You walk up and "
    "say, in plain English: “Check whether our ad spend is still paying off — and keep watch.”"
)
para("What happens next is the whole product:")
bullet([("The org decides. ", True), ("A Gemini Chief of Staff (the ADK coordinator) reasons out loud, then delegates to the right specialist — you watch the hand-off fire in real time.", False)])
bullet([("When a capability is missing, the org grows itself. ", True), ("The Chief calls the World API over MCP — world_describe → world_add_function → world_build → world_staff → world_assign. No human pre-wires a thing.", False)])
bullet([("The world rises, live. ", True), ("A new wing materialises block-by-block near HQ — and because Gemini designs each building to its function, a Payments vault looks nothing like an Analytics observatory. Every build is new.", False)])
bullet([("It staffs itself. ", True), ("A new Gemini specialist walks through the door, wired to exactly the tools that function needs.", False)])
bullet([("You walk in to live data. ", True), ("The far wall is a designed, alien-HUD dashboard of the real numbers, refreshing as the agent pulls them over MCP.", False)])
bullet([("You steer with a sentence. ", True), ("/revise show the 7-day window and flag ROAS below 2.0 — the board redraws, live.", False)])
bullet([("You stay in command. ", True), ("Every outward action — send an email, change a budget — surfaces as a tap-to-approve gate before it fires.", False)])
para(
    "You configured none of it. The org saw what it needed, built the room, hired the agent, connected "
    "the tools, and shipped — while you watched. The world is as big as the mission, and it grows itself."
)

rule()

# ── 3. Technologies used ─────────────────────────────────────────────────────
heading("Technologies used")
para(
    "The three mandatory technologies aren't decoration here — each one is load-bearing, and each maps "
    "to a file a judge can open."
)
bullet([("Gemini — the brain of every agent. ", True), ("gemini-flash-latest runs the Chief of Staff, both specialists, every hired agent, and the generative architect that designs each building. Pinned to a current GA alias (not the retired 2.0). → omo-agent/omo/agent.py, omo-agent/specialist/agent.py, runtime/src/worldArchitect.ts.", False)])
bullet([("ADK — the organisation. ", True), ("A genuine multi-agent system: an LlmAgent coordinator with sub_agents=[growth, comms] that delegates via transfer_to_agent, served by adk api_server and consumed token-by-token over /run_sse. → omo-agent/, bridged into the runtime by runtime/src/agents/AdkAgent.ts.", False)])
bullet([("MCP — every connection, and the extension mechanism itself. ", True), ("A stateless Streamable-HTTP server (official @modelcontextprotocol/sdk) exposes real business tools and the World API — so the org extends itself by calling an MCP tool, not by shipping code. ADK consumes it through McpToolset(StreamableHTTPConnectionParams). → runtime/src/mcpServer.ts.", False)])
para(
    "The rare move judges will remember: world self-extension is exposed as an MCP tool surface. The org "
    "grows the same way it does anything else — by calling a tool."
)
para(
    "Supporting stack: Paper / Java 21 plugin for the live parametric builds, villager pathing, and the "
    "tap-to-approve UI; a Node / TypeScript runtime as orchestrator; a custom client-mod for the "
    "fullscreen in-world dashboard screens."
)

rule()

# ── 4. Data sources ──────────────────────────────────────────────────────────
heading("Data sources")
para(
    "Every number on every screen is real — the agents are instructed never to invent a figure, and the "
    "architecture enforces it by giving them only live tools."
)
bullet([("Meta (Facebook) Ads — live. ", True), ("Real campaigns and performance insights (spend, impressions, clicks, CTR, CPC, CPM, reach; ROAS where available) over a selectable date window, pulled through the meta_ads_list_campaigns and meta_ads_insights MCP tools against a connected ad account. This is the primary, demo-load-bearing data source. → runtime/src/tools/metaAds.js.", False)])
bullet([("The world / org state — live. ", True), ("The organisation graph itself (HQ + every function: id, role, purpose, tools, room, staffed?) is queryable data the Chief of Staff reads via world_describe to ground itself before acting. → runtime/src/worldStore.ts.", False)])
bullet([("Live dashboard feed — derived. ", True), ("Each function's room screen binds to a real-data board (dashboard_update) that the staffing agent writes and /revise re-writes on demand.", False)])

rule()

# ── 5. Findings and learnings ────────────────────────────────────────────────
heading("Findings and learnings")
para("What the build actually taught us — the non-obvious parts:")
bullet([("Self-extension belongs in the tool layer, not the prompt. ", True), ("The breakthrough was making the World API a set of MCP tools. The moment the org grows itself the same way it sends an email — by calling a tool — extensibility stops being a developer task and becomes a sentence the user speaks. That's the innovation the architecture is built around.", False)])
bullet([("Reliability comes from constraint, not freedom. ", True), ("Free-form live building stalls on camera and free-rendered HTML looks rough. So the live beats are constrained: Gemini designs within one house style emitting a tiny, validated build-op DSL (clamped to the plot, sanitized, with a guaranteed-door post-process), and dashboards are an agent-populated design system, not agent-generated layout. The result always lands. Freedom is roadmap; constraint is what ships.", False)])
bullet([("De-risk the ADK↔Node bridge first. ", True), ("The single highest-risk seam is the SSE bridge between adk api_server and the runtime. Proving a hello-world Gemini turn over /run_sse in the first 30 minutes — with a fallback to direct Gemini OpenAI-compat — was worth more than any feature.", False)])
bullet([("Model pinning is a live wire. ", True), ("Gemini 2.0 is shut down and 2.5 deprecates mid-June 2026; pinning gemini-flash-latest is the difference between a working demo and a 404.", False)])
bullet([("The surprise: agents form a society. ", True), ("Once functions could world_consult one another — Growth asking Finance for the CAC ceiling instead of guessing — the org stopped being a star of isolated workers and started behaving like a team. Peer-to-peer consultation turned out to be the soul of the thing, not a side feature.", False)])
bullet([("Spatial oversight beats a settings page. ", True), ("Putting approvals, reasoning, and data in the world means governance is something you stand inside and veto by walking up — not a dashboard you forget to check.", False)])

rule()

# ── 6. Third-party integrations ──────────────────────────────────────────────
heading("Third-party integrations")
bullet([("Meta (Facebook) Ads API — integrated and live. ", True), ("Surfaced to the agents as MCP tools; real account data drives the demo's dashboard.", False)])
bullet([("Google Workspace MCP (Gmail / Drive) — Google's own MCP. ", True), ("The home-field integration: Gemini reasoning + ADK orchestration + the Google ecosystem reached over MCP. Outbound email is draft-only and human-gated by design — the agent prepares, you approve, then it sends.", False)])
bullet([("Stripe MCP (payments / payment links) — roadmap (v2 / north-star). ", True), ("The path to “worlds as businesses”: built worlds promoted to visitable, transactable Omo Worlds, with revenue flowing over Stripe MCP. Shown as the closing vision, not a live demo beat.", False)])
para(
    "Every integration reaches the agents the same way — over MCP — and every outward action passes "
    "through one tap-to-approve gate (120s auto-deny) before it touches the outside world. The trifecta "
    "is complete: Gemini reasons, ADK orchestrates, MCP connects — including to Google's own tools."
)

out = "/Users/harryedwards/omo-mc/docs/Omo_Mission_Control_Submission.docx"
doc.save(out)
print("wrote", out)
